import os
from datetime import datetime
from docx import Document
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape 
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT

def _fmt_rupee(value):
    try: value = int(value)
    except: return str(value)
    s = str(value)
    if len(s) <= 3: return f"₹ {s}"
    last_three = s[-3:]
    remaining = s[:-3]
    formatted_remaining = ""
    for i, digit in enumerate(reversed(remaining)):
        if i > 0 and i % 2 == 0: formatted_remaining = "," + formatted_remaining
        formatted_remaining = digit + formatted_remaining
    return f"₹ {formatted_remaining},{last_three}"

def _parse_rupee(value_str):
    try: return int(str(value_str).replace("₹", "").replace(",", "").strip())
    except: return 0

# --- WORD GENERATION (Unchanged) ---
def generate_payment_advice(subsidiary, date_str, transaction_list):
    valid_chars = "-_.() abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789"
    safe_sub_name = "".join(c for c in subsidiary if c in valid_chars).strip()
    if not os.path.exists(safe_sub_name):
        try: os.makedirs(safe_sub_name)
        except OSError as e: return False, f"Error creating folder: {e}"
    i = 1
    while True:
        filename = f"{safe_sub_name}/noting{i}.docx"
        if not os.path.exists(filename): break
        i += 1
    try:
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        p = doc.add_paragraph()
        p.add_run("Reg :-  ").bold = True
        p.runs[0].underline = True
        p.add_run("Signature of Print Payment Advice (PPA) under SDRF Parent Account").bold = True
        doc.add_paragraph()
        p2 = doc.add_paragraph()
        p2.add_run("Print payment Advice (PPA) in respect of ")
        p2.add_run(subsidiary).underline = True 
        p2.add_run(" under SDRF parent account has been placed below for JD (DM) signature please.")
        doc.add_paragraph()
        table = doc.add_table(rows=1+len(transaction_list)+1, cols=4)
        table.style = 'Table Grid'
        headers = ["Sl. No.", "In favour of", "Print Payment Advice No.", "Amount"]
        for i, t in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = t
            for p in cell.paragraphs: 
                for r in p.runs: r.font.bold = True
        total = 0
        for idx, (ppa, amt) in enumerate(transaction_list):
            row = table.rows[idx+1].cells
            total += _parse_rupee(amt)
            row[0].text = f"{idx+1}."
            row[1].text = subsidiary
            row[2].text = str(ppa)
            row[3].text = amt
        last = table.rows[-1].cells
        last[2].text = "G/TOTAL"
        last[3].text = _fmt_rupee(total)
        for c in [last[2], last[3]]:
            for p in c.paragraphs:
                for r in p.runs: r.font.bold = True
        doc.save(filename)
        return True, os.path.abspath(filename)
    except Exception as e: return False, str(e)

# --- PDF GENERATION (UPDATED) ---
def generate_summary_pdf(data_list):
    filename = f"Financial_Report_{datetime.now().strftime('%d-%m-%Y')}.pdf"
    
    try:
        # 1. Setup A4 Landscape
        doc = SimpleDocTemplate(filename, pagesize=landscape(A4), 
                                leftMargin=20, rightMargin=20, topMargin=20, bottomMargin=20)
        elements = []
        styles = getSampleStyleSheet()
        
        # 2. Title
        title_style = ParagraphStyle('CT', parent=styles['Heading1'], fontSize=16, alignment=TA_CENTER, textColor=colors.black, spaceAfter=15)
        elements.append(Paragraph("Financial Status Report (Running Balance)", title_style))
        
        date_str = datetime.now().strftime("%d-%m-%Y %H:%M %p")
        elements.append(Paragraph(f"Generated on: {date_str}", styles['Normal']))
        elements.append(Spacer(1, 15))

        # 3. Headers (2 Rows)
        # CHANGED: "Allocation Amount" -> "Previous Balance"
        row1 = [
            "Department", 
            "Previous\nBalance", 
            "Quarter 1", "", "", 
            "Quarter 2", "", "", 
            "Quarter 3", "", "", 
            "Quarter 4", "", ""
        ]
        
        row2 = [
            "", "", 
            "Addl\nAlloc", "Expenditure", "Qtr ending\nBalance",
            "Addl\nAlloc", "Expenditure", "Qtr ending\nBalance",
            "Addl\nAlloc", "Expenditure", "Qtr ending\nBalance",
            "Addl\nAlloc", "Expenditure", "Qtr ending\nBalance"
        ]

        table_data = [row1, row2]
        
        # 4. Data Rows
        for row in data_list:
            clean_row = [row[0]] # Name
            for val in row[1:]:
                txt = _fmt_rupee(val).replace("₹", "").strip() 
                clean_row.append(txt)
            table_data.append(clean_row)

        # 5. Column Widths
        cw_dept = 140
        cw_op = 60
        cw_num = 50
        col_widths = [cw_dept, cw_op] + [cw_num]*12
        
        t = Table(table_data, colWidths=col_widths, repeatRows=2)
        
        # 6. Styling
        t.setStyle(TableStyle([
            ('SPAN', (0,0), (0,1)), # Department
            ('SPAN', (1,0), (1,1)), # Previous Balance
            ('SPAN', (2,0), (4,0)), # Quarter 1
            ('SPAN', (5,0), (7,0)), # Quarter 2
            ('SPAN', (8,0), (10,0)), # Quarter 3
            ('SPAN', (11,0), (13,0)), # Quarter 4
            
            ('FONTNAME', (0,0), (-1,1), 'Helvetica-Bold'),
            ('ALIGN', (0,0), (-1,1), 'CENTER'),
            ('VALIGN', (0,0), (-1,1), 'MIDDLE'),
            ('BACKGROUND', (0,0), (-1,1), colors.lightgrey),
            ('FONTSIZE', (0,0), (-1,1), 8),
            
            ('FONTNAME', (0,2), (0,-1), 'Helvetica-Bold'),
            ('FONTSIZE', (0,2), (-1,-1), 7),
            ('ALIGN', (0,2), (0,-1), 'LEFT'),
            ('ALIGN', (1,2), (-1,-1), 'CENTER'),
            
            ('GRID', (0,0), (-1,-1), 0.5, colors.black),
            ('INNERGRID', (0,0), (-1,-1), 0.25, colors.black),
            
            ('TOPPADDING', (0,0), (-1,-1), 2),
            ('BOTTOMPADDING', (0,0), (-1,-1), 2),
            ('LEFTPADDING', (0,0), (-1,-1), 2),
            ('RIGHTPADDING', (0,0), (-1,-1), 2),
        ]))
        
        elements.append(t)
        doc.build(elements)
        return True, os.path.abspath(filename)
    except Exception as e: return False, str(e)